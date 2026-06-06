from html import escape
from pathlib import Path
from textwrap import wrap
from zipfile import ZIP_DEFLATED, ZipFile

from services.storage_service import get_chat_history, get_document

EXPORT_DIR = Path(__file__).resolve().parents[1] / "exports"

EXPORT_FORMATS = {
    "txt": "text/plain",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
EXPORT_SECTIONS = {"all", "summary", "chat", "quiz", "flashcards"}


def export_summary(document_id: str) -> Path:
    return export_document_activity(document_id, section="summary", file_format="txt")[0]


def export_chat(document_id: str) -> Path:
    return export_document_activity(document_id, section="chat", file_format="txt")[0]


def export_document_activity(document_id: str, section: str = "all", file_format: str = "txt") -> tuple[Path, str]:
    section = section.lower().strip()
    file_format = file_format.lower().strip()
    if section not in EXPORT_SECTIONS:
        raise ValueError("Pilihan aktivitas export tidak valid.")
    if file_format not in EXPORT_FORMATS:
        raise ValueError("Format export tidak valid.")

    document = get_document(document_id)
    if not document:
        raise KeyError("Dokumen tidak ditemukan.")

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    title, lines = _activity_lines(document, section)
    safe_section = section.replace("_", "-")
    path = EXPORT_DIR / f"{document_id}-{safe_section}.{file_format}"

    if file_format == "txt":
        path.write_text("\n".join([title, "", *lines]), encoding="utf-8")
    elif file_format == "pdf":
        _write_pdf(path, title, lines)
    else:
        _write_docx(path, title, lines)

    return path, EXPORT_FORMATS[file_format]


def _activity_lines(document: dict, section: str) -> tuple[str, list[str]]:
    file_name = document.get("file_name", document.get("document_id", "dokumen"))
    title = f"PDF Insight AI - {file_name}"
    sections = []

    if section in {"all", "summary"}:
        sections.append(("Ringkasan", [document.get("summary") or "Ringkasan belum tersedia."]))

    if section in {"all", "chat"}:
        chat_lines = []
        for index, item in enumerate(get_chat_history(document["document_id"]), start=1):
            chat_lines.extend([
                f"Chat {index}",
                f"Q: {item.get('question', '-')}",
                f"A: {item.get('answer', '-')}",
            ])
        sections.append(("Chat", chat_lines or ["Chat belum tersedia."]))

    if section in {"all", "quiz"}:
        quiz_items = document.get("quiz", {}).get("quiz", [])
        quiz_lines = []
        for index, item in enumerate(quiz_items, start=1):
            quiz_lines.append(f"{index}. {item.get('question', '-')}")
            for option_index, option in enumerate(item.get("options", []), start=1):
                quiz_lines.append(f"   {'ABCD'[option_index - 1]}. {option}")
            quiz_lines.append(f"   Jawaban: {item.get('answer', '-')}")
            if item.get("explanation"):
                quiz_lines.append(f"   Penjelasan: {item['explanation']}")
        sections.append(("Quiz", quiz_lines or ["Quiz belum tersedia."]))

    if section in {"all", "flashcards"}:
        cards = document.get("flashcards", {}).get("flashcards", [])
        card_lines = []
        for index, card in enumerate(cards, start=1):
            card_lines.extend([
                f"Flashcard {index}",
                f"Front: {card.get('front', '-')}",
                f"Back: {card.get('back', '-')}",
            ])
        sections.append(("Flashcards", card_lines or ["Flashcards belum tersedia."]))

    lines = []
    for heading, content in sections:
        lines.extend([heading, "-" * len(heading), *content, ""])
    return title, lines


def _write_docx(path: Path, title: str, lines: list[str]) -> None:
    paragraphs = [_docx_paragraph(title, style="Title")]
    paragraphs.extend(_docx_paragraph(line) for line in lines)
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:body>{"".join(paragraphs)}<w:sectPr/></w:body></w:document>'
    )
    with ZipFile(path, "w", ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", _docx_content_types())
        docx.writestr("_rels/.rels", _docx_relationships())
        docx.writestr("word/document.xml", document_xml)


def _docx_paragraph(text: str, style: str | None = None) -> str:
    text = escape(text or " ")
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return f"<w:p>{style_xml}<w:r><w:t xml:space=\"preserve\">{text}</w:t></w:r></w:p>"


def _docx_content_types() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )


def _docx_relationships() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '</Relationships>'
    )


def _write_pdf(path: Path, title: str, lines: list[str]) -> None:
    pdf_lines = [title, "", *lines]
    content_lines = []
    y = 800
    for line in pdf_lines:
        wrapped = wrap(line, width=88) or [""]
        for part in wrapped:
            if y < 50:
                break
            content_lines.append(f"BT /F1 10 Tf 50 {y} Td ({_pdf_escape(part)}) Tj ET")
            y -= 14
        if y < 50:
            content_lines.append(f"BT /F1 10 Tf 50 {y} Td (...konten dipotong untuk PDF sederhana) Tj ET")
            break

    stream = "\n".join(content_lines).encode("latin-1", errors="replace")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        f"5 0 obj << /Length {len(stream)} >> stream\n".encode("ascii") + stream + b"\nendstream endobj\n",
    ]
    offsets = []
    content = b"%PDF-1.4\n"
    for obj in objects:
        offsets.append(len(content))
        content += obj
    xref_offset = len(content)
    content += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii")
    for offset in offsets:
        content += f"{offset:010d} 00000 n \n".encode("ascii")
    content += f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode("ascii")
    path.write_bytes(content)


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def export_full_report_pdf(document_id: str) -> Path:
    document = get_document(document_id)
    if not document:
        raise KeyError("Dokumen tidak ditemukan.")

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    path = EXPORT_DIR / f"{document_id}-full-report.pdf"
    styles = getSampleStyleSheet()
    story = [Paragraph("PDF Insight AI - Full Analysis Report", styles["Title"]), Spacer(1, 12)]
    for heading, lines in _report_sections(document):
        story.append(Paragraph(heading, styles["Heading2"]))
        if heading == "Document Quality":
            quality = document.get("quality") or {}
            rows = [["Metric", "Value"]] + [[key.replace("_", " ").title(), str(value)] for key, value in quality.items()]
            table = Table(rows, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF2FF")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(table)
        else:
            for line in lines:
                story.append(Paragraph(_escape_reportlab(line), styles["BodyText"]))
        story.append(Spacer(1, 10))
    SimpleDocTemplate(str(path), pagesize=A4, title="PDF Insight AI Report").build(story)
    return path


def export_full_report_docx(document_id: str) -> Path:
    document = get_document(document_id)
    if not document:
        raise KeyError("Dokumen tidak ditemukan.")

    from docx import Document

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    path = EXPORT_DIR / f"{document_id}-full-report.docx"
    doc = Document()
    doc.add_heading("PDF Insight AI - Full Analysis Report", level=0)
    for heading, lines in _report_sections(document):
        doc.add_heading(heading, level=1)
        if heading == "Document Quality":
            quality = document.get("quality") or {}
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Metric"
            table.rows[0].cells[1].text = "Value"
            for key, value in quality.items():
                row = table.add_row().cells
                row[0].text = key.replace("_", " ").title()
                row[1].text = str(value)
        else:
            for line in lines:
                doc.add_paragraph(line)
    doc.save(path)
    return path


def _report_sections(document: dict) -> list[tuple[str, list[str]]]:
    analysis = document.get("analysis") or {}
    chat_history = get_chat_history(document["document_id"])
    sections = [
        ("Document Info", [
            f"File name: {document.get('file_name', '-')}",
            f"Total pages: {document.get('total_pages', '-')}",
            f"Upload date: {document.get('upload_date', '-')}",
            f"Generated date: {__import__('datetime').datetime.utcnow().isoformat()} UTC",
        ]),
        ("Document Quality", []),
        ("Summary", [document.get("summary") or analysis.get("summary") or "Ringkasan belum tersedia."]),
        ("Key Points", [f"- {item}" for item in analysis.get("key_points", [])] or ["Poin penting belum tersedia."]),
        ("Keywords", [", ".join(analysis.get("keywords", [])) or "Keywords belum tersedia."]),
        ("Suggested Questions", [f"- {item}" for item in analysis.get("suggested_questions", [])] or ["Pertanyaan saran belum tersedia."]),
        ("Chat History", _chat_report_lines(chat_history)),
        ("Quiz", _quiz_report_lines(document.get("quiz", {}).get("quiz", []))),
        ("Flashcards", _flashcard_report_lines(document.get("flashcards", {}).get("flashcards", []))),
    ]
    return sections


def _chat_report_lines(chat_history: list[dict]) -> list[str]:
    lines = []
    for index, item in enumerate(chat_history, start=1):
        lines.extend([f"Chat {index}", f"Q: {item.get('question', '-')}", f"A: {item.get('answer', '-')}"])
        for source in item.get("sources", []):
            lines.append(f"Source page {source.get('page_number', '-')}: {source.get('snippet', '-')}")
    return lines or ["Chat belum tersedia."]


def _quiz_report_lines(items: list[dict]) -> list[str]:
    lines = []
    for index, item in enumerate(items, start=1):
        lines.append(f"{index}. {item.get('question', '-')}")
        lines.append(f"Answer: {item.get('answer', '-')}")
    return lines or ["Quiz belum tersedia."]


def _flashcard_report_lines(items: list[dict]) -> list[str]:
    lines = []
    for index, item in enumerate(items, start=1):
        lines.extend([f"Flashcard {index}", f"Front: {item.get('front', '-')}", f"Back: {item.get('back', '-')}"])
    return lines or ["Flashcards belum tersedia."]


def _escape_reportlab(text: str) -> str:
    return escape(str(text)).replace("\n", "<br/>")
